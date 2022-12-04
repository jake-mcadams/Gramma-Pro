from docx import Document
import docx
import os

class ReadDoc():

    def get_paragraphs(self, doc):
        for x in doc.paragraphs:
            yield(x)
        

class CreateDoc():
    doc = docx.Document()
  
    # add a heading of level 0 (largest heading)
    doc.add_heading('Heading for the document', 0)
    
    # add a paragraph and store 
    # the object in a variable
    doc_para1 = doc.add_paragraph('Matt told her to reach for the stars, but Veronica thought it was the most ridiculous advice she\'d ever received. Sure, it had been well-meaning when he said it, but she didn\'t understand why anyone would want to suggest something that would literally kill you if you actually managed to achieve it.')
    doc_para2 = doc.add_paragraph('It was supposed to be a dream vacation. They had planned it over a year in advance so that it would be perfect in every way. It had been what they had been looking forward to through all the turmoil and negativity around them. It had been the light at the end of both their tunnels. Now that the dream vacation was only a week away, the virus had stopped all air travel.')
    doc_para3 = doc.add_paragraph('A long black shadow slid across the pavement near their feet and the five Venusians, very much startled, looked overhead. They were barely in time to see the huge gray form of the carnivore before it vanished behind a sign atop a nearby building which bore the mystifying information "Pepsi-Cola."')
    
    # add a run i.e, style like 
    # bold, italic, underline, etc.
    # doc_para.add_run('hey there, bold here').bold = True
    # doc_para.add_run(', and ')
    # doc_para.add_run('these words are italic').italic = True
    
    # add a page break to start a new page
    doc.add_page_break()
    
    # add a heading of level 2
    doc.add_heading('Heading level 2', 2)
    
    # now save the document to a location
    doc.save('./testing.docx')
    

class Doc(ReadDoc, CreateDoc):
    def __init__(self, doc) -> None:
        self.doc = doc
    
    def get_doc_path(self):
        file_path = os.path.basename(self.doc)
        return file_path
    
    
    
test = Doc('../testing.docx')
print(test.get_doc_path())
print(test.get_paragraphs())